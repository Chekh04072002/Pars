import sys
from PyQt5 import QtCore, QtGui, QtWidgets
from main05 import Ui_MainWindow
from glob import glob
import os
import docx

app = QtWidgets.QApplication(sys.argv)

MainWindow = QtWidgets.QMainWindow()
ui = Ui_MainWindow()
ui.setupUi(MainWindow)

MainWindow.show()

yesList_obrazovanie = ['', ' ', 'Начальное', 'начальное', 'Среднее', 'среднее', 'Высшее', 'высшее']
yesList_stepen = ['', ' ', 'Кандидат', 'кандидат', 'Доктор', 'доктор', 'Нет', 'Нет']
yesList_opit = ['', ' ', ]
for i in range(100):
    yesList_opit.append(str(i))

def sort():
    os.mkdir("Sorted")
    obrazovanie = ui.lineEdit_1.text()
    stepen = ui.lineEdit_2.text()
    opit = ui.lineEdit_3.text()
    opit = opit
    print(obrazovanie)
    print(stepen)
    print(opit)
    if obrazovanie not in yesList_obrazovanie or stepen not in yesList_stepen or opit not in yesList_opit:
        if obrazovanie not in yesList_obrazovanie:
            ui.lineEdit.setText('Введите требования об образовании так, как указано в скобках')
        if stepen not in yesList_stepen:
            ui.lineEdit.setText('Введите требования о научной степени так, как указано в скобках')
        if opit not in yesList_opit:
            ui.lineEdit.setText('Введите требования к опыту так, как указано в скобках')
    else:
        #  - - -
        if obrazovanie == '' and stepen == '' and opit == '':
            ui.lineEdit.setText('Введите требования к сортировке')

        # + - -
        elif obrazovanie != '' and stepen == '' and opit == '':
            for i in glob('*.docx'):  # В директории ищет нужный документ
                doc = docx.Document(i)
                text = []
                for paragraph in doc.paragraphs:  # Считываю тект по параграфам
                    text.append(paragraph.text)
                yestext = '\n'.join(text)
                if obrazovanie in yestext:
                    os.rename(f"C://Users/user/Desktop/Parse/{i}", f"C://Users/user/Desktop/Parse/Sorted/{i}")
                else:
                    ui.lineEdit.setText('Сортировка закончена, больше нет совпадений')

        # - + -
        elif obrazovanie == '' and stepen != '' and opit == '':
            for i in glob('*.docx'):  # В директории ищет нужный документ
                doc = docx.Document(i)
                text = []
                for paragraph in doc.paragraphs:  # Считываю тект по параграфам
                    text.append(paragraph.text)
                yestext = '\n'.join(text)
                if stepen in yestext:
                    os.rename(f"C://Users/user/Desktop/Parse/{i}", f"C://Users/user/Desktop/Parse/Sorted/{i}")
                else:
                    ui.lineEdit.setText('Сортировка закончена, больше нет совпадений')

        # - - +
        elif obrazovanie == '' and stepen == '' and opit != '':
            for i in glob('*.docx'):  # В директории ищет нужный документ
                doc = docx.Document(i)
                text = []
                for paragraph in doc.paragraphs:  # Считываю тект по параграфам
                    text.append(paragraph.text)
                yestext = '\n'.join(text)
                for q in range(100):
                    if str(q) in yestext:
                        op = q
                if int(opit) <= int(op):
                    os.rename(f"C://Users/user/Desktop/Parse/{i}", f"C://Users/user/Desktop/Parse/Sorted/{i}")
                else:
                    ui.lineEdit.setText('Сортировка закончена, больше нет совпадений')

        # + + -
        elif obrazovanie != '' and stepen != '' and opit == '':
            for i in glob('*.docx'):  # В директории ищет нужный документ
                doc = docx.Document(i)
                text = []
                for paragraph in doc.paragraphs:  # Считываю тект по параграфам
                    text.append(paragraph.text)
                yestext = '\n'.join(text)
                if obrazovanie in yestext and stepen in yestext:
                    os.rename(f"C://Users/user/Desktop/Parse/{i}", f"C://Users/user/Desktop/Parse/Sorted/{i}")
                else:
                    ui.lineEdit.setText('Сортировка закончена, больше нет совпадений')

        # + - +
        elif obrazovanie != '' and stepen == '' and opit != '':
            for i in glob('*.docx'):  # В директории ищет нужный документ
                doc = docx.Document(i)
                text = []
                for paragraph in doc.paragraphs:  # Считываю тект по параграфам
                    text.append(paragraph.text)
                yestext = '\n'.join(text)
                for q in range(100):
                    if str(q) in yestext:
                        op = q
                if int(opit) <= int(op) and obrazovanie in yestext:
                    os.rename(f"C://Users/user/Desktop/Parse/{i}", f"C://Users/user/Desktop/Parse/Sorted/{i}")
                else:
                    ui.lineEdit.setText('Сортировка закончена, больше нет совпадений')

        # - + +
        elif obrazovanie == '' and stepen != '' and opit != '':
            for i in glob('*.docx'):  # В директории ищет нужный документ
                doc = docx.Document(i)
                text = []
                for paragraph in doc.paragraphs:  # Считываю тект по параграфам
                    text.append(paragraph.text)
                yestext = '\n'.join(text)
                for q in range(100):
                    if str(q) in yestext:
                        op = q
                if int(opit) <= int(op) and stepen in yestext:
                    os.rename(f"C://Users/user/Desktop/Parse/{i}", f"C://Users/user/Desktop/Parse/Sorted/{i}")
                else:
                    ui.lineEdit.setText('Сортировка закончена, больше нет совпадений')

        # + + +
        elif obrazovanie != '' and stepen != '' and opit != '':
            for i in glob('*.docx'):  # В директории ищет нужный документ
                doc = docx.Document(i)
                text = []
                for paragraph in doc.paragraphs:  # Считываю тект по параграфам
                    text.append(paragraph.text)
                yestext = '\n'.join(text)
                for q in range(100):
                    if str(q) in yestext:
                        op = q
                if int(opit) <= int(op) and stepen in yestext and obrazovanie in yestext:
                    os.rename(f"C://Users/user/Desktop/Parse/{i}", f"C://Users/user/Desktop/Parse/Sorted/{i}")
                else:
                    ui.lineEdit.setText('Сортировка закончена, больше нет совпадений')



ui.pushButton.clicked.connect(sort)

sys.exit(app.exec_())
